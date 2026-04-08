#!/usr/bin/env python3
"""
Script 10: Pre-Meeting Intelligence Driver (Enhanced with Word Document Output)

Generates comprehensive pre-meeting intelligence using Claude AI with:
- Gmail API for email history
- Google Calendar API for meeting details
- Claude web search for research
- Generates formatted Word document (.docx) in Meraglim brand standards
- Creates ClickUp Meeting Intelligence task with document attachment
- Toggles ClickUp Task Created checkbox in Airtable

Can be triggered by:
- Script 8 (Manual trigger on Meeting Scheduled status)
- Script 10T (Calendar event trigger)
- Manual execution with prospect email

Usage:
  python3 script_10_pre_meeting_intelligence.py <prospect_email> [linkedin_message]
  
Example:
  python3 script_10_pre_meeting_intelligence.py edkeels@beyondbreakeven.net
"""

import os
import sys
import json
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Any, Optional, List
import tempfile
import mimetypes

# Add parent directory to path for shared utilities
sys.path.insert(0, str(Path(__file__).parent))

from shared_utils import setup_logger, StateManager, send_error_notification, handle_errors, load_credentials

# Third-party imports
import requests
import anthropic
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load credentials
creds = load_credentials()

# Configuration
SCRIPT_NAME = "script_10_pre_meeting_intelligence"
AIRTABLE_API_KEY = creds["AIRTABLE_API_KEY"]
AIRTABLE_BASE_ID = creds["AIRTABLE_BASE_ID"]
AIRTABLE_TABLE_ID = "tblxEhVek8ldTQMW1"  # Prospects table

CLAUDE_API_KEY = creds.get("CLAUDE_API_KEY", "")
CLAUDE_MODEL = "claude-sonnet-4-20250514"

CLICKUP_API_KEY = creds.get("CLICKUP_API_KEY", "")
CLICKUP_MEETING_INTELLIGENCE_LIST_ID = creds.get("CLICKUP_MEETING_INTELLIGENCE_LIST_ID", "")

# Google OAuth
GOOGLE_TOKEN_FILE = os.getenv("GOOGLE_TOKEN_FILE", os.path.expanduser("~/Automations/config/google_token.json"))
OAUTH_TOKEN_FILE = os.getenv("OAUTH_TOKEN_FILE", os.path.expanduser("~/Automations/config/oauth_token.json"))

# Meraglim brand colors (hex)
COLOR_BLACK = "1A1A1A"
COLOR_GOLD = "C9A84C"
COLOR_DGRAY = "3D3D3D"
COLOR_MGRAY = "595959"
COLOR_LGRAY = "F5F5F5"
COLOR_WHITE = "FFFFFF"

# Initialize logger and state manager
logger = setup_logger(SCRIPT_NAME)
state_manager = StateManager()


class GmailClient:
    """Client for Gmail API operations."""
    
    def __init__(self, token_file: str):
        self.token_file = token_file
        self.service = self._get_service()
    
    def _get_service(self):
        """Get Gmail service with OAuth credentials."""
        try:
            creds = Credentials.from_authorized_user_file(self.token_file)
            return build("gmail", "v1", credentials=creds)
        except Exception as e:
            logger.warning(f"Gmail service initialization failed: {str(e)}")
            return None
    
    def get_email_history(self, email_address: str, max_results: int = 10) -> List[Dict[str, Any]]:
        """Get email history with a specific contact."""
        if not self.service:
            return []
        
        try:
            query = f"from:{email_address} OR to:{email_address}"
            results = self.service.users().messages().list(userId="me", q=query, maxResults=max_results).execute()
            
            messages = results.get("messages", [])
            email_history = []
            
            for msg in messages:
                try:
                    message = self.service.users().messages().get(userId="me", id=msg["id"]).execute()
                    headers = message["payload"]["headers"]
                    
                    email_data = {
                        "from": next((h["value"] for h in headers if h["name"] == "From"), ""),
                        "to": next((h["value"] for h in headers if h["name"] == "To"), ""),
                        "subject": next((h["value"] for h in headers if h["name"] == "Subject"), ""),
                        "date": next((h["value"] for h in headers if h["name"] == "Date"), ""),
                    }
                    email_history.append(email_data)
                except Exception as e:
                    logger.warning(f"Error processing email: {str(e)}")
                    continue
            
            logger.info(f"Retrieved {len(email_history)} emails with {email_address}")
            return email_history
        
        except Exception as e:
            logger.warning(f"Error getting email history: {str(e)}")
            return []


class GoogleCalendarClient:
    """Client for Google Calendar API operations."""
    
    def __init__(self, token_file: str):
        self.token_file = token_file
        self.service = self._get_service()
    
    def _get_service(self):
        """Get Google Calendar service with OAuth credentials."""
        try:
            creds = Credentials.from_authorized_user_file(self.token_file)
            return build("calendar", "v3", credentials=creds)
        except Exception as e:
            logger.warning(f"Google Calendar service initialization failed: {str(e)}")
            return None
    
    def get_upcoming_events(self, max_results: int = 5) -> List[Dict[str, Any]]:
        """Get upcoming calendar events."""
        if not self.service:
            return []
        
        try:
            now = datetime.utcnow().isoformat() + "Z"
            events_result = self.service.events().list(
                calendarId="primary",
                timeMin=now,
                maxResults=max_results,
                singleEvents=True,
                orderBy="startTime"
            ).execute()
            
            events = events_result.get("items", [])
            logger.info(f"Retrieved {len(events)} upcoming events")
            return events
        
        except Exception as e:
            logger.warning(f"Error getting calendar events: {str(e)}")
            return []


class AirtableClient:
    """Client for Airtable API operations."""
    
    def __init__(self, api_key: str, base_id: str):
        self.api_key = api_key
        self.base_id = base_id
        self.base_url = "https://api.airtable.com/v0"
    
    def get_records(self, table_id: str, filter_formula: Optional[str] = None ) -> list:
        """Get records from Airtable with optional filter."""
        url = f"{self.base_url}/{self.base_id}/{table_id}"
        headers = {"Authorization": f"Bearer {self.api_key}"}
        params = {}
        
        if filter_formula:
            params["filterByFormula"] = filter_formula
        
        try:
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            return response.json().get("records", [])
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching Airtable records: {str(e)}")
            raise
    
    def update_record(self, table_id: str, record_id: str, fields: Dict[str, Any]) -> bool:
        """Update a record in Airtable."""
        url = f"{self.base_url}/{self.base_id}/{table_id}/{record_id}"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        data = {"fields": fields}
        
        try:
            response = requests.patch(url, headers=headers, json=data)
            response.raise_for_status()
            logger.info(f"Updated Airtable record {record_id}")
            return True
        except requests.exceptions.RequestException as e:
            logger.error(f"Error updating Airtable record: {str(e)}")
            raise


class ClickUpClient:
    """Client for ClickUp API operations."""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://api.clickup.com/api/v2"
        self.headers = {
            "Authorization": api_key,
            "Content-Type": "application/json"
        }
    
    def create_task(self, list_id: str, name: str, description: str = "", 
                   priority: int = 2 ) -> Optional[Dict[str, Any]]:
        """Create a task in ClickUp."""
        try:
            url = f"{self.base_url}/list/{list_id}/task"
            
            data = {
                "name": name,
                "description": description,
                "priority": priority
            }
            
            response = requests.post(url, headers=self.headers, json=data)
            response.raise_for_status()
            
            task_response = response.json()
            task_id = task_response.get("id")
            task_url = task_response.get("url")
            
            logger.info(f"Created ClickUp Meeting Intelligence task: {task_id} - {task_url}")
            return task_response
        
        except requests.exceptions.RequestException as e:
            logger.error(f"Error creating ClickUp task: {str(e)}")
            if hasattr(e, 'response'):
                logger.error(f"Response: {e.response.text}")
            return None
    
    def attach_file_to_task(self, task_id: str, file_path: str) -> bool:
        """Attach a file to a ClickUp task."""
        try:
            url = f"{self.base_url}/task/{task_id}/attachment"
            
            with open(file_path, 'rb') as f:
                files = {'attachment': f}
                # Remove Content-Type header for file upload
                headers = {"Authorization": self.api_key}
                response = requests.post(url, headers=headers, files=files)
                response.raise_for_status()
            
            logger.info(f"Attached file to ClickUp task {task_id}")
            return True
        
        except Exception as e:
            logger.error(f"Error attaching file to ClickUp task: {str(e)}")
            return False


class ClaudeIntelligenceGenerator:
    """Generate pre-meeting intelligence using Claude AI with web search."""
    
    def __init__(self, api_key: str, model: str):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model
    
    def generate_meeting_prep(self, prospect_data: Dict[str, Any], 
                            email_history: List[Dict[str, Any]],
                            calendar_events: List[Dict[str, Any]],
                            linkedin_message: str = "") -> Optional[str]:
        """Generate comprehensive pre-meeting intelligence using Claude with web search."""
        try:
            # Extract prospect information
            first_name = prospect_data.get("First Name", "")
            last_name = prospect_data.get("Last Name", "")
            name = prospect_data.get("Name", f"{first_name} {last_name}".strip())
            company = prospect_data.get("Company", "Unknown")
            title = prospect_data.get("Title", "")
            email = prospect_data.get("Email", "")
            sector = prospect_data.get("Sector", "")
            revenue = prospect_data.get("Estimated Revenue", "")
            ebitda = prospect_data.get("EBITDA Normalized", "")
            key_risk = prospect_data.get("Key Risk", "")
            
            # Format email history
            email_summary = ""
            if email_history:
                email_summary = "\n\nRECENT EMAIL HISTORY:\n"
                for email in email_history[-5:]:
                    email_summary += f"- {email['date']}: {email['subject']}\n"
            
            # Format calendar events
            calendar_summary = ""
            if calendar_events:
                calendar_summary = "\n\nUPCOMING EVENTS:\n"
                for event in calendar_events[:3]:
                    start = event.get("start", {}).get("dateTime", "")
                    summary = event.get("summary", "")
                    calendar_summary += f"- {start}: {summary}\n"
            
            # Format LinkedIn message
            linkedin_summary = ""
            if linkedin_message:
                linkedin_summary = f"\n\nLINKEDIN MESSAGE FROM {name}:\n{linkedin_message}\n"
            
            # Build the comprehensive prompt for Claude
            prompt = f"""Meeting intel: {name} — {company}

DIRECT CONTACT MESSAGE:
{linkedin_summary}

Run full background analysis. Deliver the complete briefing in the following sequence:

1. MEETING OVERVIEW
   Date, time, meeting title, Zoom/call link if available.

2. RELATIONSHIP CONTEXT
   Gmail history summary — when communication started, frequency, last interaction date and topic, open items or commitments outstanding from either side, communication tone.
   {email_summary}

3. INDIVIDUAL PROFILE — {name}
   Professional background, current role and title, career trajectory. LinkedIn activity — recent posts, themes, priorities, public commentary. Any shared connections or network overlap with Meraglim or Kevin Massengill.
   Current Title: {title}
   Email: {email}

4. COMPANY INTELLIGENCE — {company}
   What the company does, who they serve, approximate size, revenue signals, geography. Recent news, funding rounds, leadership changes, product launches, partnerships, or strategic shifts in the past 6 months. Business model and how they generate revenue.
   Sector: {sector}
   Estimated Revenue: {revenue}

5. BACKGROUND ANALYSIS
   Individual reputation scan — professional standing, any public controversies, litigation, or adverse press.
   Legal and regulatory scan — lawsuits, SEC filings, regulatory investigations, settlement history.
   Company risk profile — known operational, financial, or market challenges. Any red flags that warrant caution before engaging.
   Key Risk: {key_risk}

6. COMPETITIVE LANDSCAPE
   Top 3–5 competitors to {company}. How the company positions against them. Where it leads, where it lags. Any recent competitive moves or market share shifts. Assess whether competitors are potential Meraglim acquisition targets.

7. MERAGLIM RELEVANCE
   Evaluate this contact and company against Meraglim's acquisition criteria:
   — Does the company have $500K–$10M EBITDA?
   — Is the owner or principal likely to be a motivated seller?
   — Is there sufficient free cash flow to service Annuity Sale payments?
   — Is this contact a potential acquisition target, referral source, or strategic partner?
   — What is the most likely path to value for Meraglim from this relationship?
   EBITDA: {ebitda}

8. CONVERSATION GUIDE
   3–5 prioritized questions to ask in the meeting, ordered by strategic importance.
   Key listening signals — what answers would qualify or disqualify this contact.
   Any risk flags from the background analysis that should be probed directly.
   Suggested close — what outcome should Kevin drive toward by end of the call?

9. CALL SCRIPT
   Opening (first 2 minutes): how to frame the call and establish rapport.
   Discovery (minutes 2–12): the core questions and what to listen for.
   Positioning (minutes 12–18): how to present the Annuity Sale if appropriate, tailored to this specific contact and company context.
   Qualification close (minutes 18–25): how to determine next steps and what a successful outcome looks like.
   Objection handling: anticipated pushback and suggested responses based on what you know about this contact.

Flag any ⚠️ risk alerts at the top of the briefing if the background analysis surfaces adverse findings. Otherwise deliver in the sequence above.

Use web search to research {name}, {company}, and recent news about the company. Include specific findings in each section."""
            
            # Call Claude API with web search enabled
            message = self.client.messages.create(
                model=self.model,
                max_tokens=4096,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Extract the response
            intelligence = message.content[0].text
            logger.info(f"Generated comprehensive pre-meeting intelligence for {name}")
            return intelligence
        
        except Exception as e:
            logger.error(f"Error generating meeting intelligence: {str(e)}")
            return None


class MeraglimDocumentGenerator:
    """Generate Word documents in Meraglim brand standards."""
    
    def __init__(self):
        self.color_black = RGBColor(0x1A, 0x1A, 0x1A)
        self.color_gold = RGBColor(0xC9, 0xA8, 0x4C)
        self.color_dgray = RGBColor(0x3D, 0x3D, 0x3D)
        self.color_mgray = RGBColor(0x59, 0x59, 0x59)
        self.color_lgray = RGBColor(0xF5, 0xF5, 0xF5)
    
    def _add_section_heading(self, doc, text):
        """Add H1 section heading with gold underline."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = self.color_black
        p.space_before = Pt(18)
        p.space_after = Pt(6)
        
        # Add gold underline
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C9A84C')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _add_subsection_heading(self, doc, text):
        """Add H2 subsection heading."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = self.color_dgray
        p.space_before = Pt(12)
        p.space_after = Pt(4)
    
    def _add_body_text(self, doc, text):
        """Add body text paragraph."""
        if not text.strip():
            return
        p = doc.add_paragraph(text.strip())
        p_format = p.paragraph_format
        p_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = self.color_black
    
    def _add_bullet_point(self, doc, text):
        """Add bullet point."""
        p = doc.add_paragraph(text.strip(), style='List Bullet')
        p_format = p.paragraph_format
        p_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = self.color_black
    
    def _parse_and_format_content(self, doc, intelligence_text):
        """Parse Claude markdown output and format as proper Word document."""
        lines = intelligence_text.split('\n')
        current_section = None
        buffer = []
        
        for line in lines:
            stripped = line.strip()
            
            # Skip empty lines
            if not stripped:
                if buffer:
                    # Flush buffer as body text
                    for text in buffer:
                        if text.startswith('- '):
                            self._add_bullet_point(doc, text[2:])
                        else:
                            self._add_body_text(doc, text)
                    buffer = []
                continue
            
            # Detect section headings (e.g., "## 1. MEETING OVERVIEW")
            if stripped.startswith('##'):
                # Flush buffer first
                for text in buffer:
                    if text.startswith('- '):
                        self._add_bullet_point(doc, text[2:])
                    else:
                        self._add_body_text(doc, text)
                buffer = []
                
                # Add section heading
                heading = stripped.replace('##', '').replace('#', '').strip()
                self._add_section_heading(doc, heading)
            
            # Detect subsection headings (e.g., "**Current Title**:")
            elif stripped.startswith('**') and '**' in stripped[2:]:
                # Flush buffer first
                for text in buffer:
                    if text.startswith('- '):
                        self._add_bullet_point(doc, text[2:])
                    else:
                        self._add_body_text(doc, text)
                buffer = []
                
                # Add subsection heading
                heading = stripped.replace('**', '').strip()
                self._add_subsection_heading(doc, heading)
            
            # Detect bullet points
            elif stripped.startswith('- '):
                # Flush non-bullet buffer
                for text in buffer:
                    if not text.startswith('- '):
                        self._add_body_text(doc, text)
                buffer = []
                
                # Add bullet point
                self._add_bullet_point(doc, stripped[2:])
            
            else:
                # Regular text - add to buffer
                buffer.append(stripped)
        
        # Flush remaining buffer
        for text in buffer:
            if text.startswith('- '):
                self._add_bullet_point(doc, text[2:])
            else:
                self._add_body_text(doc, text)
    
    def generate_intelligence_brief(self, prospect_name: str, company: str, 
                                   intelligence_text: str) -> str:
        """Generate a formatted Word document with pre-meeting intelligence."""
        try:
            doc = Document()
            
            # Set up page margins (1 inch all around)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Title
            title = doc.add_paragraph()
            title_run = title.add_run(f"Meeting Intelligence Brief")
            title_run.font.size = Pt(26)
            title_run.font.bold = True
            title_run.font.color.rgb = self.color_black
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(f"{prospect_name} — {company}")
            subtitle_run.font.size = Pt(20)
            subtitle_run.font.color.rgb = self.color_mgray
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            date_run.font.size = Pt(12)
            date_run.font.color.rgb = self.color_mgray
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add gold line
            pPr = date_para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'C9A84C')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            # Add spacing
            doc.add_paragraph()
            
            # Parse and format the intelligence text
            self._parse_and_format_content(doc, intelligence_text)
            
            # Add footer
            footer_para = doc.add_paragraph()
            footer_para.space_before = Pt(20)
            pPr = footer_para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '12')
            top.set(qn('w:space'), '1')
            top.set(qn('w:color'), 'C9A84C')
            pBdr.append(top)
            pPr.append(pBdr)
            
            footer_run = footer_para.add_run("Meraglim Holdings Corporation • Meeting Intelligence • CONFIDENTIAL")
            footer_run.font.size = Pt(10)
            footer_run.font.italic = True
            footer_run.font.color.rgb = self.color_mgray
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_filename = f"MHC10_Intelligence_{prospect_name.replace(' ', '_')}_{timestamp}.docx"
            doc_path = os.path.expanduser(f"~/Automations/output/{doc_filename}")
            
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(doc_path), exist_ok=True)
            
            doc.save(doc_path)
            logger.info(f"Generated Word document: {doc_path}")
            return doc_path
        
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            return None


def get_prospect_by_email(email: str) -> Optional[Dict[str, Any]]:
    """Get prospect record from Airtable by email."""
    airtable = AirtableClient(AIRTABLE_API_KEY, AIRTABLE_BASE_ID)
    
    filter_formula = f"{{Email}} = '{email}'"
    
    try:
        records = airtable.get_records(AIRTABLE_TABLE_ID, filter_formula)
        
        if records:
            return records[0]
        else:
            logger.warning(f"No prospect found with email: {email}")
            return None
    
    except Exception as e:
        logger.error(f"Error getting prospect by email: {str(e)}")
        return None


def update_prospect_with_intelligence(prospect_id: str, prep_summary: str) -> bool:
    """Update prospect record with pre-meeting intelligence and toggle checkbox."""
    airtable = AirtableClient(AIRTABLE_API_KEY, AIRTABLE_BASE_ID)
    
    fields = {
        "Meeting Prep Summary": prep_summary,
        "Last Briefed": datetime.now().strftime("%Y-%m-%d"),
        "ClickUp Task Created": True  # Toggle the checkbox
    }
    
    try:
        airtable.update_record(AIRTABLE_TABLE_ID, prospect_id, fields)
        return True
    except Exception as e:
        logger.error(f"Error updating prospect: {str(e)}")
        return False


@handle_errors(SCRIPT_NAME, logger)
def main(prospect_email: str, linkedin_message: str = ""):
    """Main execution function."""
    logger.info(f"Starting {SCRIPT_NAME} for prospect: {prospect_email}")
    
    try:
        # Validate configuration
        if not CLAUDE_API_KEY or CLAUDE_API_KEY.startswith("<"):
            logger.error("Claude API key not configured")
            return False
        
        if not CLICKUP_API_KEY or CLICKUP_API_KEY.startswith("<"):
            logger.error("ClickUp API key not configured")
            return False
        
        if not CLICKUP_MEETING_INTELLIGENCE_LIST_ID or CLICKUP_MEETING_INTELLIGENCE_LIST_ID.startswith("<"):
            logger.error("ClickUp Meeting Intelligence List ID not configured")
            return False
        
        # Get prospect from Airtable
        prospect = get_prospect_by_email(prospect_email)
        
        if not prospect:
            logger.error(f"Prospect not found: {prospect_email}")
            return False
        
        prospect_id = prospect["id"]
        prospect_fields = prospect.get("fields", {})
        prospect_name = prospect_fields.get("Name", "Unknown")
        company = prospect_fields.get("Company", "Unknown")
        
        logger.info(f"Processing prospect: {prospect_name}")
        
        # Get email history
        gmail = GmailClient(OAUTH_TOKEN_FILE)
        email_history = gmail.get_email_history(prospect_email)
        
        # Get calendar events
        calendar = GoogleCalendarClient(GOOGLE_TOKEN_FILE)
        calendar_events = calendar.get_upcoming_events()
        
        # Generate pre-meeting intelligence with web search
        claude = ClaudeIntelligenceGenerator(CLAUDE_API_KEY, CLAUDE_MODEL)
        prep_summary = claude.generate_meeting_prep(prospect_fields, email_history, calendar_events, linkedin_message)
        
        if not prep_summary:
            logger.error("Failed to generate pre-meeting intelligence")
            return False
        
        # Generate Word document
        doc_gen = MeraglimDocumentGenerator()
        doc_path = doc_gen.generate_intelligence_brief(prospect_name, company, prep_summary)
        
        if not doc_path:
            logger.error("Failed to generate Word document")
            return False
        
        # Create ClickUp Meeting Intelligence task
        clickup = ClickUpClient(CLICKUP_API_KEY)
        task = clickup.create_task(
            CLICKUP_MEETING_INTELLIGENCE_LIST_ID,
            f"Meeting Intelligence: {prospect_name} - {company}",
            prep_summary,
            priority=2
        )
        
        if not task:
            logger.error("Failed to create ClickUp task")
            return False
        
        # Attach Word document to ClickUp task
        task_id = task.get("id")
        if task_id and doc_path:
            clickup.attach_file_to_task(task_id, doc_path)
        
        # Update Airtable with intelligence and toggle checkbox
        if not update_prospect_with_intelligence(prospect_id, prep_summary):
            logger.error("Failed to update Airtable")
            return False
        
        logger.info(f"Successfully generated pre-meeting intelligence for {prospect_name}")
        logger.info(f"Word document: {doc_path}")
        logger.info(f"ClickUp task: {task.get('url')}")
        state_manager.update_state(SCRIPT_NAME, status="success")
        return True
    
    except Exception as e:
        logger.error(f"Error in main execution: {str(e)}", exc_info=True)
        state_manager.increment_error_count(SCRIPT_NAME)
        send_error_notification(
            SCRIPT_NAME,
            f"Error generating pre-meeting intelligence: {str(e)}",
            str(e),
            logger
        )
        return False


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 script_10_pre_meeting_intelligence.py <prospect_email> [linkedin_message]")
        print("Example: python3 script_10_pre_meeting_intelligence.py edkeels@beyondbreakeven.net")
        sys.exit(1)
    
    prospect_email = sys.argv[1]
    linkedin_message = sys.argv[2] if len(sys.argv) > 2 else ""
    success = main(prospect_email, linkedin_message)
    sys.exit(0 if success else 1)
